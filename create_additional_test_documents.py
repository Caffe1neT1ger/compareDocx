"""
Скрипт для создания дополнительных масштабных тестовых DOCX документов.
Документы с разными стилями для тестирования сравнения по содержимому.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw
import os
import re


def create_simple_image(width, height, color, text, filename):
    """Создание простого тестового изображения."""
    img = Image.new('RGB', (width, height), color=color)
    draw = ImageDraw.Draw(img)
    
    # Рамка
    draw.rectangle([10, 10, width-10, height-10], outline='black', width=2)
    
    # Текст
    if text:
        text_y = height // 2 - 10
        draw.text((width // 2 - 50, text_y), text, fill='black')
    
    img.save(filename)
    return filename


def create_additional_document_1():
    """Создание первого дополнительного документа с кастомными стилями."""
    doc = Document()
    
    # Заголовок документа (без использования стандартного стиля Heading)
    title_para = doc.add_paragraph('Проект разработки корпоративной информационной системы')
    title_run = title_para.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 128)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1 - используем обычный стиль с форматированием
    doc.add_paragraph()  # Пустая строка
    
    section1 = doc.add_paragraph('1. Введение и общие сведения')
    section1.runs[0].font.size = Pt(16)
    section1.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Данный проект направлен на создание комплексной информационной системы '
        'для управления бизнес-процессами крупного предприятия. Система должна '
        'обеспечивать интеграцию всех подразделений и автоматизацию ключевых операций.'
    )
    
    # Подраздел 1.1 - используем обычный стиль
    sub1 = doc.add_paragraph('1.1. Цели и задачи проекта')
    sub1.runs[0].font.size = Pt(14)
    sub1.runs[0].font.bold = True
    sub1.runs[0].font.italic = True
    
    doc.add_paragraph(
        'Основными целями проекта являются повышение эффективности работы предприятия, '
        'снижение операционных затрат и улучшение качества обслуживания клиентов.'
    )
    
    # Таблица 1 - Цели проекта
    doc.add_paragraph()  # Пустая строка
    
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Цель'
    header_cells[1].text = 'Описание'
    
    # Заголовок таблицы - жирный
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data1 = [
        ['Автоматизация', 'Автоматизация основных бизнес-процессов'],
        ['Интеграция', 'Интеграция всех подразделений'],
        ['Аналитика', 'Внедрение системы аналитики и отчетности'],
        ['Оптимизация', 'Оптимизация рабочих процессов']
    ]
    
    for i, row_data in enumerate(data1, 1):
        row_cells = table1.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Раздел 2
    doc.add_page_break()
    
    section2 = doc.add_paragraph('2. Технические требования')
    section2.runs[0].font.size = Pt(16)
    section2.runs[0].font.bold = True
    section2.runs[0].font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph(
        'Система должна соответствовать следующим техническим требованиям:'
    )
    
    # Подраздел 2.1
    sub2_1 = doc.add_paragraph('2.1. Требования к производительности')
    sub2_1.runs[0].font.size = Pt(14)
    sub2_1.runs[0].font.bold = True
    
    doc.add_paragraph('• Обработка не менее 10000 транзакций в минуту', style='List Bullet')
    doc.add_paragraph('• Поддержка до 2000 одновременных пользователей', style='List Bullet')
    doc.add_paragraph('• Время отклика не более 1 секунды', style='List Bullet')
    
    # Таблица 2 - Технические характеристики
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Light Grid Accent 1'
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Минимум'
    header_cells[2].text = 'Рекомендуется'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 128)
    
    data2 = [
        ['Процессор', '8 ядер', '16 ядер'],
        ['ОЗУ', '32 ГБ', '64 ГБ'],
        ['Диск', '1 ТБ SSD', '2 ТБ SSD'],
        ['Сеть', '10 Гбит/с', '25 Гбит/с'],
        ['ОС', 'Linux/Windows', 'Linux/Windows']
    ]
    
    for i, row_data in enumerate(data2, 1):
        row_cells = table2.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 1
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 2.1 - Архитектура системы:')
    img1_path = create_simple_image(500, 300, 'lightblue', 'Архитектура v3.0', 'documents/temp_img_add1_1.png')
    doc.add_picture(img1_path, width=Inches(5))
    
    # Раздел 3
    doc.add_page_break()
    
    section3 = doc.add_paragraph('3. Функциональные требования')
    section3.runs[0].font.size = Pt(16)
    section3.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Система должна обеспечивать выполнение следующих функций:'
    )
    
    # Подраздел 3.1
    sub3_1 = doc.add_paragraph('3.1. Модуль управления пользователями')
    sub3_1.runs[0].font.size = Pt(14)
    sub3_1.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Модуль должен обеспечивать создание, редактирование и удаление учетных записей. '
        'Поддержка ролевой модели доступа и многофакторной аутентификации.'
    )
    
    # Таблица 3 - Роли пользователей
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Light Grid Accent 1'
    header_cells = table3.rows[0].cells
    header_cells[0].text = 'Роль'
    header_cells[1].text = 'Права доступа'
    header_cells[2].text = 'Описание'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data3 = [
        ['Администратор', 'Полный', 'Полный доступ ко всем функциям'],
        ['Менеджер', 'Расширенный', 'Доступ к управлению и отчетам'],
        ['Пользователь', 'Базовый', 'Доступ к основным функциям'],
        ['Гость', 'Ограниченный', 'Только просмотр'],
        ['Аудитор', 'Чтение', 'Доступ только для чтения']
    ]
    
    for i, row_data in enumerate(data3, 1):
        row_cells = table3.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 2
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 3.1 - Схема модулей:')
    img2_path = create_simple_image(450, 250, 'lightgreen', 'Модули системы', 'documents/temp_img_add1_2.png')
    doc.add_picture(img2_path, width=Inches(4.5))
    
    # Раздел 4
    doc.add_page_break()
    
    section4 = doc.add_paragraph('4. Этапы реализации')
    section4.runs[0].font.size = Pt(16)
    section4.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Проект будет реализован в несколько этапов:'
    )
    
    # Таблица 4 - Этапы проекта
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Light Grid Accent 1'
    header_cells = table4.rows[0].cells
    header_cells[0].text = 'Этап'
    header_cells[1].text = 'Срок'
    header_cells[2].text = 'Результат'
    header_cells[3].text = 'Ответственный'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data4 = [
        ['Анализ', '1 месяц', 'Техническое задание', 'Аналитики'],
        ['Проектирование', '2 месяца', 'Проектная документация', 'Архитекторы'],
        ['Разработка', '6 месяцев', 'Рабочая версия', 'Разработчики'],
        ['Тестирование', '2 месяца', 'Протестированная система', 'Тестировщики'],
        ['Внедрение', '1 месяц', 'Внедренная система', 'Внедренцы']
    ]
    
    for i, row_data in enumerate(data4, 1):
        row_cells = table4.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Заключение
    doc.add_page_break()
    
    conclusion = doc.add_paragraph('Заключение')
    conclusion.runs[0].font.size = Pt(16)
    conclusion.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Реализация проекта позволит значительно повысить эффективность работы предприятия '
        'и улучшить качество обслуживания клиентов. Срок реализации проекта составляет '
        '12 месяцев с момента начала работ.'
    )
    
    # Удаление временных файлов
    temp_files = [
        'documents/temp_img_add1_1.png',
        'documents/temp_img_add1_2.png'
    ]
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)
    
    doc.save('documents/additional_test_document_1.docx')
    print("Создан файл: documents/additional_test_document_1.docx")


def create_additional_document_2():
    """Создание второго дополнительного документа (с изменениями и разными стилями)."""
    doc = Document()
    
    # Заголовок документа (другой стиль)
    title_para = doc.add_paragraph('Проект разработки корпоративной информационной системы')
    title_run = title_para.runs[0]
    title_run.font.size = Pt(22)  # Изменен размер
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 200)  # Изменен цвет
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1 - другой стиль форматирования
    doc.add_paragraph()  # Пустая строка
    
    section1 = doc.add_paragraph('1. Введение и общие сведения')
    section1.runs[0].font.size = Pt(18)  # Изменен размер
    section1.runs[0].font.bold = True
    section1.runs[0].font.underline = True  # Добавлено подчеркивание
    
    # Грамматические изменения
    doc.add_paragraph(
        'Данный проект направлен на создание комплексной информационной системы '
        'для управления бизнес-процессами крупного предприятия. Система должна '
        'обеспечивать интеграцию всех подразделений, автоматизацию ключевых операций, и повышение эффективности.'  # Добавлен текст
    )
    
    # Подраздел 1.1 - другой стиль
    sub1 = doc.add_paragraph('1.1. Цели и задачи проекта')
    sub1.runs[0].font.size = Pt(15)  # Изменен размер
    sub1.runs[0].font.bold = True
    # Убрано italic
    
    doc.add_paragraph(
        'Основными целями проекта являются повышение эффективности работы предприятия, '
        'снижение операционных затрат, улучшение качества обслуживания клиентов, и внедрение инновационных технологий.'  # Добавлен текст
    )
    
    # Таблица 1 - Цели проекта (с изменениями)
    doc.add_paragraph()  # Пустая строка
    
    table1 = doc.add_table(rows=6, cols=2)  # Добавлена строка
    table1.style = 'Light Grid Accent 1'
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Цель'
    header_cells[1].text = 'Описание'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Изменен цвет
    
    data1 = [
        ['Автоматизация', 'Автоматизация основных бизнес-процессов и рабочих операций'],  # Изменено
        ['Интеграция', 'Интеграция всех подразделений и внешних систем'],  # Изменено
        ['Аналитика', 'Внедрение системы аналитики, отчетности и бизнес-интеллекта'],  # Изменено
        ['Оптимизация', 'Оптимизация рабочих процессов и ресурсов'],  # Изменено
        ['Безопасность', 'Обеспечение информационной безопасности']  # Новая строка
    ]
    
    for i, row_data in enumerate(data1, 1):
        row_cells = table1.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Раздел 2 - другой стиль
    doc.add_page_break()
    
    section2 = doc.add_paragraph('2. Технические требования')
    section2.runs[0].font.size = Pt(18)  # Изменен размер
    section2.runs[0].font.bold = True
    section2.runs[0].font.color.rgb = RGBColor(200, 0, 0)  # Изменен цвет
    
    doc.add_paragraph(
        'Система должна соответствовать следующим техническим требованиям:'
    )
    
    # Подраздел 2.1
    sub2_1 = doc.add_paragraph('2.1. Требования к производительности')
    sub2_1.runs[0].font.size = Pt(15)  # Изменен размер
    sub2_1.runs[0].font.bold = True
    
    doc.add_paragraph('• Обработка не менее 15000 транзакций в минуту', style='List Bullet')  # Изменено
    doc.add_paragraph('• Поддержка до 3000 одновременных пользователей', style='List Bullet')  # Изменено
    doc.add_paragraph('• Время отклика не более 0.8 секунды', style='List Bullet')  # Изменено
    doc.add_paragraph('• Доступность системы 99.95%', style='List Bullet')  # Новая строка
    
    # Таблица 2 - Технические характеристики (с изменениями)
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Light Grid Accent 1'
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Минимум'
    header_cells[2].text = 'Рекомендуется'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 200)  # Изменен цвет
    
    data2 = [
        ['Процессор', '12 ядер', '24 ядра'],  # Изменено
        ['ОЗУ', '64 ГБ', '128 ГБ'],  # Изменено
        ['Диск', '2 ТБ SSD', '4 ТБ SSD'],  # Изменено
        ['Сеть', '25 Гбит/с', '40 Гбит/с'],  # Изменено
        ['ОС', 'Linux/Windows Server', 'Linux/Windows Server']
    ]
    
    for i, row_data in enumerate(data2, 1):
        row_cells = table2.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 1 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 2.1 - Архитектура системы:')
    img1_path = create_simple_image(500, 300, 'lightgreen', 'Архитектура v3.1', 'documents/temp_img_add2_1.png')  # Изменен цвет и текст
    doc.add_picture(img1_path, width=Inches(5))
    
    # Раздел 3
    doc.add_page_break()
    
    section3 = doc.add_paragraph('3. Функциональные требования')
    section3.runs[0].font.size = Pt(18)  # Изменен размер
    section3.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Система должна обеспечивать выполнение следующих функций:'
    )
    
    # Подраздел 3.1
    sub3_1 = doc.add_paragraph('3.1. Модуль управления пользователями')
    sub3_1.runs[0].font.size = Pt(15)  # Изменен размер
    sub3_1.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Модуль должен обеспечивать создание, редактирование, и удаление учетных записей. '  # Добавлена запятая
        'Поддержка ролевой модели доступа, многофакторной аутентификации, и интеграции с Active Directory.'  # Добавлен текст
    )
    
    # Таблица 3 - Роли пользователей (с изменениями)
    table3 = doc.add_table(rows=7, cols=3)  # Добавлена строка
    table3.style = 'Light Grid Accent 1'
    header_cells = table3.rows[0].cells
    header_cells[0].text = 'Роль'
    header_cells[1].text = 'Права доступа'
    header_cells[2].text = 'Описание'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data3 = [
        ['Администратор', 'Полный', 'Полный доступ ко всем функциям системы'],
        ['Менеджер', 'Расширенный', 'Доступ к управлению, отчетам, и аналитике'],  # Изменено
        ['Пользователь', 'Базовый', 'Доступ к основным функциям и личным данным'],  # Изменено
        ['Гость', 'Ограниченный', 'Только просмотр публичной информации'],  # Изменено
        ['Аудитор', 'Чтение', 'Доступ только для чтения и аудита'],  # Изменено
        ['Модератор', 'Модерация', 'Доступ к модерации контента']  # Новая строка
    ]
    
    for i, row_data in enumerate(data3, 1):
        row_cells = table3.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 2 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 3.1 - Схема модулей:')
    img2_path = create_simple_image(450, 250, 'lightyellow', 'Модули системы v2', 'documents/temp_img_add2_2.png')  # Изменен цвет и текст
    doc.add_picture(img2_path, width=Inches(4.5))
    
    # Новое изображение
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 3.2 - Диаграмма взаимодействия:')
    img3_path = create_simple_image(400, 200, 'lightcoral', 'Взаимодействие', 'documents/temp_img_add2_3.png')
    doc.add_picture(img3_path, width=Inches(4))
    
    # Раздел 4
    doc.add_page_break()
    
    section4 = doc.add_paragraph('4. Этапы реализации')
    section4.runs[0].font.size = Pt(18)  # Изменен размер
    section4.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Проект будет реализован в несколько этапов:'
    )
    
    # Таблица 4 - Этапы проекта (с изменениями)
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Light Grid Accent 1'
    header_cells = table4.rows[0].cells
    header_cells[0].text = 'Этап'
    header_cells[1].text = 'Срок'
    header_cells[2].text = 'Результат'
    header_cells[3].text = 'Ответственный'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data4 = [
        ['Анализ', '1.5 месяца', 'Техническое задание и анализ требований', 'Аналитики'],  # Изменено
        ['Проектирование', '2.5 месяца', 'Проектная документация и архитектура', 'Архитекторы'],  # Изменено
        ['Разработка', '8 месяцев', 'Рабочая версия системы', 'Разработчики'],  # Изменено
        ['Тестирование', '2 месяца', 'Протестированная и оптимизированная система', 'Тестировщики'],  # Изменено
        ['Внедрение', '1.5 месяца', 'Внедренная и настроенная система', 'Внедренцы']  # Изменено
    ]
    
    for i, row_data in enumerate(data4, 1):
        row_cells = table4.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Заключение
    doc.add_page_break()
    
    conclusion = doc.add_paragraph('Заключение')
    conclusion.runs[0].font.size = Pt(18)  # Изменен размер
    conclusion.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Реализация проекта позволит значительно повысить эффективность работы предприятия, '
        'улучшить качество обслуживания клиентов, и внедрить современные технологии управления. '  # Добавлен текст
        'Срок реализации проекта составляет 15 месяцев с момента начала работ.'  # Изменено
    )
    
    # Удаление временных файлов
    temp_files = [
        'documents/temp_img_add2_1.png',
        'documents/temp_img_add2_2.png',
        'documents/temp_img_add2_3.png'
    ]
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)
    
    doc.save('documents/additional_test_document_2.docx')
    print("Создан файл: documents/additional_test_document_2.docx")


if __name__ == "__main__":
    print("Создание дополнительных масштабных тестовых документов...")
    print("=" * 70)
    
    create_additional_document_1()
    create_additional_document_2()
    
    print("=" * 70)
    print("Все дополнительные тестовые документы созданы успешно!")
    print("\nСозданные файлы:")
    print("  - documents/additional_test_document_1.docx (базовый документ с кастомными стилями)")
    print("  - documents/additional_test_document_2.docx (версия с изменениями и разными стилями)")
    print("\nОсобенности документов:")
    print("  - Использование кастомных стилей (не стандартные Heading)")
    print("  - Разное форматирование (размеры шрифтов, цвета, подчеркивание)")
    print("  - 4 раздела с подразделами")
    print("  - 4 таблицы в каждом документе")
    print("  - 2-3 изображения в каждом документе")
    print("  - Грамматические изменения (запятые, пунктуация)")
    print("  - Изменения в таблицах и изображениях")
    print("  - Идеально для тестирования сравнения по содержимому!")

