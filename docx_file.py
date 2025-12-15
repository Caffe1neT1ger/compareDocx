"""
Модуль для работы с DOCX файлами.
Класс DocxFile предоставляет функционал для парсинга и анализа структуры документа.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional, Tuple
import re
import hashlib
import io
from config import config
from logger_config import logger
from exceptions import DocumentLoadError, DocumentParseError, ValidationError
from validators import validate_file_size, validate_document_structure

# Импорт tqdm и colorama для прогресс-баров (опционально)
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    def tqdm(iterable, *args, **kwargs):
        return iterable

try:
    from colorama import init, Fore, Style
    init(autoreset=True)
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False
    class Fore:
        BLUE = ''
        CYAN = ''
        GREEN = ''
        YELLOW = ''
        MAGENTA = ''
        RESET = ''
    class Style:
        RESET_ALL = ''


class DocxFile:
    """Класс для работы с DOCX файлами."""
    
    def __init__(self, file_path: str):
        """
        Инициализация класса.
        
        Args:
            file_path: Путь к DOCX файлу
        
        Raises:
            DocumentLoadError: Если не удалось загрузить документ
            DocumentParseError: Если произошла ошибка при парсинге
        """
        self.file_path = file_path
        
        # Валидация размера файла
        try:
            from pathlib import Path
            validate_file_size(Path(file_path))
        except Exception as e:
            logger.error(f"Ошибка валидации файла {file_path}: {e}")
            raise DocumentLoadError(file_path, str(e))
        
        # Загрузка документа
        try:
            self.document = Document(file_path)
            logger.debug(f"Документ загружен: {file_path}")
        except Exception as e:
            logger.error(f"Ошибка загрузки документа {file_path}: {e}")
            raise DocumentLoadError(file_path, str(e))
        
        self.paragraphs = []
        self.sections = []
        self.chapters = []
        self.tables = []
        self.images = []
        self.hierarchy_stack = []  # Стек для построения иерархии
        
        # Парсинг документа
        try:
            self._parse_document()
            logger.info(
                f"Документ распарсен: {len(self.paragraphs)} абзацев, "
                f"{len(self.tables)} таблиц, {len(self.images)} изображений"
            )
        except Exception as e:
            logger.error(f"Ошибка парсинга документа {file_path}: {e}")
            raise DocumentParseError(file_path, str(e))
    
    def _parse_document(self):
        """Парсинг документа: извлечение абзацев, разделов, глав, таблиц и изображений."""
        current_section = None
        current_chapter = None
        section_index = 0
        chapter_index = 0
        paragraph_index = 0
        chars_per_page = config.document.chars_per_page
        
        # Парсинг таблиц
        self._parse_tables()
        
        # Парсинг изображений
        self._parse_images()
        
        # Валидация структуры документа
        try:
            validate_document_structure(
                len(self.document.paragraphs),
                len(self.tables),
                len(self.images)
            )
        except ValidationError as e:
            logger.warning(f"Предупреждение о структуре документа: {e}")
            # Продолжаем работу, но логируем предупреждение
        
        # Оптимизация: вычисляем накопленную длину символов один раз
        total_chars = 0
        
        # Парсинг абзацев с прогресс-баром
        paragraphs_list = list(self.document.paragraphs)
        paragraphs_iter = paragraphs_list
        if TQDM_AVAILABLE and len(paragraphs_list) > 5:  # Показываем прогресс только для больших документов
            if COLORAMA_AVAILABLE:
                color_desc = f"{Fore.CYAN}{Style.BRIGHT}Парсинг абзацев{Style.RESET_ALL}"
            else:
                color_desc = "Парсинг абзацев"
            paragraphs_iter = tqdm(paragraphs_list, desc=color_desc, unit="абзац", leave=False, ncols=100,
                                   bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}]')
        
        for para in paragraphs_iter:
            text = para.text.strip()
            
            # Обновление иерархии при встрече заголовка
            if text:
                style_name = para.style.name if para.style else "Normal"
                heading_level = self._get_heading_level(style_name, para)
                element_type = self._classify_element(text, heading_level, style_name)
                
                # Игнорируем абзацы, которые являются подписями к рисункам/таблицам
                # (не добавляем их в иерархию)
                if not (text.lower().startswith('рисунок') or 
                        text.lower().startswith('таблица') or
                        text.lower().startswith('рис.') or
                        text.lower().startswith('табл.')):
                    if element_type in ["section", "chapter"]:
                        # Обновление стека иерархии
                        self._update_hierarchy_stack(text, heading_level, element_type)
            
            if not text:
                continue
            
            paragraph_index += 1
            
            # Определение стиля абзаца
            style_name = para.style.name if para.style else "Normal"
            
            # Определение уровня заголовка
            heading_level = self._get_heading_level(style_name, para)
            
            # Определение типа элемента
            element_type = self._classify_element(text, heading_level, style_name)
            
            if element_type == "section":
                section_index += 1
                current_section = {
                    "index": section_index,
                    "title": text,
                    "level": heading_level,
                    "paragraphs": []
                }
                self.sections.append(current_section)
                current_chapter = None
                chapter_index = 0
            
            elif element_type == "chapter":
                chapter_index += 1
                current_chapter = {
                    "index": chapter_index,
                    "title": text,
                    "level": heading_level,
                    "paragraphs": []
                }
                if current_section:
                    if "chapters" not in current_section:
                        current_section["chapters"] = []
                    current_section["chapters"].append(current_chapter)
                else:
                    self.chapters.append(current_chapter)
            
            # Оптимизация: вычисляем страницу на основе накопленной длины
            total_chars += len(text)
            estimated_page = max(1, (total_chars // chars_per_page) + 1)
            
            # Построение полного пути
            full_path = self._build_full_path(element_type)
            
            # Добавление абзаца
            paragraph_data = {
                "text": text,
                "style": style_name,
                "level": heading_level,
                "type": element_type,
                "section_index": section_index if current_section else None,
                "chapter_index": chapter_index if current_chapter else None,
                "paragraph_index": paragraph_index,
                "full_path": full_path,
                "page": estimated_page
            }
            
            self.paragraphs.append(paragraph_data)
            
            # Добавление в текущий раздел/главу
            if current_chapter:
                current_chapter["paragraphs"].append(paragraph_data)
            elif current_section:
                current_section["paragraphs"].append(paragraph_data)
    
    def _get_heading_level(self, style_name: str, para) -> int:
        """
        Определение уровня заголовка.
        Поддерживает как стандартные стили Word, так и кастомные стили.
        
        Args:
            style_name: Название стиля
            para: Объект абзаца
            
        Returns:
            Уровень заголовка (0 - обычный текст, 1-9 - уровни заголовков)
        """
        # 1. Проверка outline level стиля (работает для кастомных стилей тоже)
        # Это самый надежный способ, так как работает с любыми стилями, у которых задан outline level
        try:
            if para.style:
                # Проверяем outline level через paragraph_format
                if hasattr(para.style, 'paragraph_format'):
                    pf = para.style.paragraph_format
                    if hasattr(pf, 'outline_level') and pf.outline_level is not None:
                        level = pf.outline_level
                        if level >= 1 and level <= 9:
                            return level
                
                # Альтернативный способ через XML (для кастомных стилей)
                try:
                    style_xml = para.style._element
                    if style_xml is not None:
                        # Ищем outline level в XML
                        outline_elem = style_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                        if outline_elem is not None and outline_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') is not None:
                            level = int(outline_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                            if level >= 0 and level <= 8:  # В XML уровни 0-8 соответствуют заголовкам 1-9
                                return level + 1
                except (AttributeError, ValueError, TypeError):
                    pass
        except (AttributeError, TypeError):
            pass
        
        # 2. Проверка стандартных стилей заголовков Word
        if "Heading" in style_name:
            match = re.search(r'Heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))
            return 1
        
        # 3. Проверка кастомных стилей, которые могут быть основаны на заголовках
        # Ищем паттерны типа "Заголовок 1", "Заголовок_1", "H1", "Title", "Subtitle" и т.д.
        heading_patterns = [
            (r'заголовок\s*(\d+)', True),  # Заголовок 1, Заголовок_1
            (r'h\s*(\d+)', True),  # H1, H 1
            (r'title', 1),  # Title
            (r'subtitle', 2),  # Subtitle
            (r'heading\s*(\d+)', True),  # heading 1 (разные регистры)
        ]
        
        style_lower = style_name.lower()
        for pattern, level in heading_patterns:
            if isinstance(level, bool) and level:
                match = re.search(pattern, style_lower)
                if match:
                    return int(match.group(1))
            elif isinstance(level, int):
                if re.search(pattern, style_lower):
                    return level
        
        # 4. Анализ форматирования текста (эвристика для кастомных стилей)
        if para.runs:
            first_run = para.runs[0]
            text = para.text.strip()
            
            # Проверка размера шрифта (заголовки обычно крупнее)
            font_size = None
            try:
                if first_run.font and first_run.font.size:
                    font_size = first_run.font.size.pt
            except:
                pass
            
            # Проверка форматирования
            is_bold = first_run.bold
            is_italic = first_run.italic
            
            # Эвристика: короткий жирный текст с большим шрифтом - вероятно заголовок
            if is_bold and len(text) < 150:
                # Если шрифт крупнее 14pt - вероятно заголовок уровня 1-2
                if font_size and font_size >= 14:
                    if font_size >= 18:
                        return 1
                    elif font_size >= 16:
                        return 2
                    else:
                        return 3
                # Если нет информации о размере, но текст короткий и жирный
                elif len(text) < 50:
                    return 1
                elif len(text) < 100:
                    return 2
        
        # 5. Проверка по структуре текста (нумерованные заголовки)
        text = para.text.strip()
        # Паттерн: "1. ", "1.1. ", "1.1.1. " и т.д.
        numbered_pattern = re.match(r'^(\d+(?:\.\d+)*)\.\s+', text)
        if numbered_pattern:
            numbers = numbered_pattern.group(1).split('.')
            level = len(numbers)
            if level <= 9:
                return level
        
        return 0
    
    def _classify_element(self, text: str, heading_level: int, style_name: str) -> str:
        """
        Классификация элемента документа.
        
        Args:
            text: Текст элемента
            heading_level: Уровень заголовка
            style_name: Название стиля
            
        Returns:
            Тип элемента: "section", "chapter", "paragraph"
        """
        if heading_level >= 1:
            # Заголовки уровня 1-2 считаем разделами
            if heading_level <= 2:
                return "section"
            # Заголовки уровня 3+ считаем главами
            else:
                return "chapter"
        
        # Проверка на заголовок по формату (нумерованные заголовки)
        if re.match(r'^\d+\.?\s+[А-ЯЁA-Z]', text) or re.match(r'^[А-ЯЁA-Z][а-яёa-z]+\s+\d+', text):
            if len(text) < 100:
                return "section"
        
        return "paragraph"
    
    def get_all_paragraphs(self) -> List[Dict]:
        """
        Получить все абзацы документа.
        
        Returns:
            Список словарей с данными абзацев
        """
        return self.paragraphs
    
    def get_sections(self) -> List[Dict]:
        """
        Получить все разделы документа.
        
        Returns:
            Список разделов
        """
        return self.sections
    
    def get_chapters(self) -> List[Dict]:
        """
        Получить все главы документа.
        
        Returns:
            Список глав
        """
        return self.chapters
    
    def get_paragraphs_by_section(self, section_index: int) -> List[Dict]:
        """
        Получить абзацы конкретного раздела.
        
        Args:
            section_index: Индекс раздела
            
        Returns:
            Список абзацев раздела
        """
        if 0 <= section_index < len(self.sections):
            return self.sections[section_index].get("paragraphs", [])
        return []
    
    def get_paragraphs_by_chapter(self, chapter_index: int) -> List[Dict]:
        """
        Получить абзацы конкретной главы.
        
        Args:
            chapter_index: Индекс главы
            
        Returns:
            Список абзацев главы
        """
        if 0 <= chapter_index < len(self.chapters):
            return self.chapters[chapter_index].get("paragraphs", [])
        return []
    
    def _update_hierarchy_stack(self, text: str, level: int, element_type: str):
        """
        Обновление стека иерархии при встрече заголовка.
        
        Args:
            text: Текст заголовка
            level: Уровень заголовка
            element_type: Тип элемента
        """
        # Удаляем элементы с уровнем >= текущего
        while self.hierarchy_stack and self.hierarchy_stack[-1]["level"] >= level:
            self.hierarchy_stack.pop()
        
        # Добавляем новый элемент
        self.hierarchy_stack.append({
            "text": text,
            "level": level,
            "type": element_type
        })
    
    def _build_full_path(self, element_type: str) -> str:
        """
        Построение полного пути до элемента.
        
        Args:
            element_type: Тип элемента
            
        Returns:
            Полный путь в формате "Раздел 1 > Подраздел 1.1 > Пункт 1.1.1"
        """
        if not self.hierarchy_stack:
            return ""
        
        path_parts = []
        for item in self.hierarchy_stack:
            text = item["text"]
            level = item["level"]
            
            # Извлечение номера из текста (если есть)
            number_match = re.match(r'^(\d+(?:\.\d+)*)', text)
            if number_match:
                number = number_match.group(1)
                # Определение типа по уровню и номеру
                if level <= 2:
                    path_parts.append(f"Раздел {number}")
                elif level <= 4:
                    path_parts.append(f"Подраздел {number}")
                else:
                    path_parts.append(f"Пункт {number}")
            else:
                # Если номера нет, используем текст
                short_text = text[:50] if len(text) > 50 else text
                if level <= 2:
                    path_parts.append(f"Раздел: {short_text}")
                elif level <= 4:
                    path_parts.append(f"Подраздел: {short_text}")
                else:
                    path_parts.append(f"Пункт: {short_text}")
        
        return " > ".join(path_parts) if path_parts else ""
    
    def _parse_tables(self):
        """Парсинг таблиц из документа."""
        tables_list = list(enumerate(self.document.tables))
        tables_iter = tables_list
        if TQDM_AVAILABLE and len(tables_list) > 0:
            if COLORAMA_AVAILABLE:
                color_desc = f"{Fore.YELLOW}{Style.BRIGHT}Парсинг таблиц{Style.RESET_ALL}"
            else:
                color_desc = "Парсинг таблиц"
            tables_iter = tqdm(tables_list, total=len(tables_list), desc=color_desc, unit="таблица", leave=False, ncols=100,
                              bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}]')
        
        for table_idx, table in tables_iter:
            table_data = {
                "index": table_idx + 1,
                "rows": [],
                "row_count": len(table.rows),
                "col_count": len(table.columns) if table.rows else 0
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data["rows"].append(row_data)
            
            # Создание текстового представления таблицы для сравнения
            table_text = "\n".join(["\t".join(row) for row in table_data["rows"]])
            table_data["text"] = table_text
            table_data["hash"] = hashlib.md5(table_text.encode()).hexdigest()
            
            self.tables.append(table_data)
    
    def _parse_images(self):
        """Парсинг изображений из документа."""
        image_idx = 0
        processed_rels = set()
        
        # Поиск изображений через relationships
        try:
            for rel in self.document.part.rels.values():
                if rel.target_ref and "image" in str(rel.target_ref).lower():
                    rel_id = getattr(rel, 'rId', None)
                    if rel_id and rel_id not in processed_rels:
                        processed_rels.add(rel_id)
                        
                        image_data = {
                            "index": image_idx + 1,
                            "type": "image",
                            "relationship_id": rel_id
                        }
                        
                        # Попытка получить данные изображения
                        try:
                            if hasattr(rel, 'target_part'):
                                image_part = rel.target_part
                                if hasattr(image_part, 'blob'):
                                    image_data["size"] = len(image_part.blob)
                                    image_data["hash"] = hashlib.md5(image_part.blob).hexdigest()
                                else:
                                    image_data["hash"] = f"img_{rel_id}_{image_idx}"
                            else:
                                image_data["hash"] = f"img_{rel_id}_{image_idx}"
                        except Exception:
                            image_data["hash"] = f"img_{rel_id}_{image_idx}"
                        
                        self.images.append(image_data)
                        image_idx += 1
        except Exception:
            # Если не удалось извлечь через relationships, пробуем через XML
            try:
                for para_idx, para in enumerate(self.document.paragraphs):
                    for run in para.runs:
                        xml_str = str(run._element.xml)
                        if 'pic:pic' in xml_str or 'w:drawing' in xml_str:
                            image_data = {
                                "index": image_idx + 1,
                                "paragraph_index": para_idx + 1,
                                "type": "image",
                                "hash": f"image_{para_idx}_{image_idx}"
                            }
                            self.images.append(image_data)
                            image_idx += 1
                            break
            except Exception:
                pass
    
    def get_tables(self) -> List[Dict]:
        """
        Получить все таблицы документа.
        
        Returns:
            Список таблиц
        """
        return self.tables
    
    def get_images(self) -> List[Dict]:
        """
        Получить все изображения документа.
        
        Returns:
            Список изображений
        """
        return self.images
    
    def get_structure_info(self) -> Dict:
        """
        Получить информацию о структуре документа.
        
        Returns:
            Словарь с информацией о структуре
        """
        return {
            "total_paragraphs": len(self.paragraphs),
            "total_sections": len(self.sections),
            "total_chapters": len(self.chapters),
            "total_tables": len(self.tables),
            "total_images": len(self.images),
            "sections": [
                {
                    "index": s["index"],
                    "title": s["title"],
                    "paragraphs_count": len(s.get("paragraphs", [])),
                    "chapters_count": len(s.get("chapters", []))
                }
                for s in self.sections
            ]
        }

