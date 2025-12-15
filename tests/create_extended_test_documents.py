"""
Скрипт для создания масштабных тестовых DOCX документов для полноценного тестирования.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import os
import random


def create_simple_image(width, height, color, text, filename):
    """Создание простого тестового изображения."""
    img = Image.new('RGB', (width, height), color=color)
    draw = ImageDraw.Draw(img)
    
    # Рамка
    draw.rectangle([10, 10, width-10, height-10], outline='black', width=2)
    
    # Текст
    if text:
        # Простой текст по центру
        text_y = height // 2 - 10
        draw.text((width // 2 - 50, text_y), text, fill='black')
    
    img.save(filename)
    return filename


def create_extended_document_1():
    """Создание первого масштабного тестового документа."""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Техническое задание на разработку информационной системы управления предприятием', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Содержание
    doc.add_page_break()
    doc.add_heading('Содержание', 1)
    doc.add_paragraph('1. Общие положения')
    doc.add_paragraph('2. Требования к системе')
    doc.add_paragraph('3. Функциональные требования')
    doc.add_paragraph('4. Технические требования')
    doc.add_paragraph('5. Требования к интерфейсу')
    doc.add_paragraph('6. Требования к безопасности')
    doc.add_paragraph('7. Этапы разработки')
    doc.add_paragraph('8. Заключение')
    
    # Раздел 1
    doc.add_page_break()
    doc.add_heading('1. Общие положения', 1)
    
    doc.add_paragraph(
        'Настоящее техническое задание определяет требования к разработке '
        'информационной системы управления предприятием. Система предназначена '
        'для комплексной автоматизации всех бизнес-процессов организации.'
    )
    
    doc.add_heading('1.1. Назначение системы', 2)
    
    doc.add_paragraph(
        'Информационная система управления предприятием предназначена для решения '
        'следующих основных задач:'
    )
    
    doc.add_paragraph('• Автоматизация учета и управления ресурсами предприятия', style='List Bullet')
    doc.add_paragraph('• Управление документооборотом и делопроизводством', style='List Bullet')
    doc.add_paragraph('• Контроль выполнения задач и проектов', style='List Bullet')
    doc.add_paragraph('• Аналитика и формирование отчетности', style='List Bullet')
    doc.add_paragraph('• Интеграция с внешними системами', style='List Bullet')
    
    doc.add_heading('1.2. Область применения', 2)
    
    doc.add_paragraph(
        'Система предназначена для использования в организациях различных отраслей '
        'экономики. Может быть адаптирована под специфику конкретного предприятия.'
    )
    
    # Таблица 1 - Области применения
    doc.add_heading('1.3. Области применения системы', 2)
    
    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Light Grid Accent 1'
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Отрасль'
    header_cells[1].text = 'Применение'
    
    data1 = [
        ['Производство', 'Управление производственными процессами'],
        ['Торговля', 'Управление продажами и складом'],
        ['Услуги', 'Управление клиентской базой'],
        ['Строительство', 'Управление проектами'],
        ['Образование', 'Управление учебным процессом']
    ]
    
    for i, row_data in enumerate(data1, 1):
        row_cells = table1.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 1
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 1.1 - Схема архитектуры системы:')
    img1_path = create_simple_image(500, 300, 'lightblue', 'Архитектура v1.0', '../documents/temp_img_ext1_1.png')
    doc.add_picture(img1_path, width=Inches(5))
    
    # Раздел 2
    doc.add_page_break()
    doc.add_heading('2. Требования к системе', 1)
    
    doc.add_heading('2.1. Общие требования', 2)
    
    doc.add_paragraph(
        'Система должна обеспечивать надежную работу в режиме 24/7. '
        'Время отклика системы не должно превышать 2 секунд для стандартных операций.'
    )
    
    doc.add_heading('2.2. Требования к производительности', 2)
    
    doc.add_paragraph(
        'Система должна обрабатывать не менее 5000 транзакций в минуту. '
        'Поддерживать одновременную работу не менее 500 пользователей.'
    )
    
    # Таблица 2 - Требования к производительности
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Light Grid Accent 1'
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Требование'
    header_cells[2].text = 'Единица измерения'
    
    data2 = [
        ['Пропускная способность', '5000', 'транзакций/мин'],
        ['Время отклика', '2', 'секунд'],
        ['Количество пользователей', '500', 'одновременно'],
        ['Доступность', '99.9', '%']
    ]
    
    for i, row_data in enumerate(data2, 1):
        row_cells = table2.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('2.3. Требования к масштабируемости', 2)
    
    doc.add_paragraph(
        'Система должна поддерживать горизонтальное масштабирование. '
        'Возможность добавления новых серверов без остановки работы системы.'
    )
    
    # Изображение 2
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 2.1 - Схема масштабирования:')
    img2_path = create_simple_image(400, 250, 'lightgreen', 'Масштабирование', '../documents/temp_img_ext1_2.png')
    doc.add_picture(img2_path, width=Inches(4))
    
    # Раздел 3
    doc.add_page_break()
    doc.add_heading('3. Функциональные требования', 1)
    
    doc.add_heading('3.1. Модуль управления пользователями', 2)
    
    doc.add_paragraph(
        'Модуль должен обеспечивать создание, редактирование и удаление учетных записей пользователей. '
        'Поддержка ролевой модели доступа.'
    )
    
    doc.add_heading('3.1.1. Управление ролями', 3)
    
    doc.add_paragraph('Система должна поддерживать следующие роли:')
    doc.add_paragraph('• Администратор системы', style='List Bullet')
    doc.add_paragraph('• Менеджер', style='List Bullet')
    doc.add_paragraph('• Пользователь', style='List Bullet')
    doc.add_paragraph('• Гость', style='List Bullet')
    
    doc.add_heading('3.1.2. Аутентификация', 3)
    
    doc.add_paragraph(
        'Система должна обеспечивать безопасную аутентификацию пользователей. '
        'Поддержка двухфакторной аутентификации.'
    )
    
    # Таблица 3 - Роли и права доступа
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Light Grid Accent 1'
    header_cells = table3.rows[0].cells
    header_cells[0].text = 'Роль'
    header_cells[1].text = 'Просмотр'
    header_cells[2].text = 'Редактирование'
    header_cells[3].text = 'Удаление'
    
    data3 = [
        ['Администратор', 'Да', 'Да', 'Да'],
        ['Менеджер', 'Да', 'Да', 'Нет'],
        ['Пользователь', 'Да', 'Нет', 'Нет'],
        ['Гость', 'Ограничен', 'Нет', 'Нет']
    ]
    
    for i, row_data in enumerate(data3, 1):
        row_cells = table3.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('3.2. Модуль управления документами', 2)
    
    doc.add_paragraph(
        'Модуль обеспечивает полный цикл работы с документами: создание, редактирование, '
        'согласование, утверждение и архивирование.'
    )
    
    doc.add_heading('3.2.1. Создание документов', 3)
    
    doc.add_paragraph(
        'Система должна поддерживать создание документов различных типов: '
        'приказы, распоряжения, договоры, акты и другие.'
    )
    
    doc.add_heading('3.2.2. Маршрутизация документов', 3)
    
    doc.add_paragraph(
        'Система должна обеспечивать настройку маршрутов движения документов. '
        'Поддержка параллельного и последовательного согласования.'
    )
    
    # Изображение 3
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 3.1 - Схема маршрутизации документов:')
    img3_path = create_simple_image(450, 300, 'lightyellow', 'Маршрутизация', '../documents/temp_img_ext1_3.png')
    doc.add_picture(img3_path, width=Inches(4.5))
    
    doc.add_heading('3.3. Модуль отчетности', 2)
    
    doc.add_paragraph(
        'Модуль должен обеспечивать формирование различных типов отчетов: '
        'статистические, аналитические, регламентированные.'
    )
    
    # Таблица 4 - Типы отчетов
    table4 = doc.add_table(rows=6, cols=2)
    table4.style = 'Light Grid Accent 1'
    header_cells = table4.rows[0].cells
    header_cells[0].text = 'Тип отчета'
    header_cells[1].text = 'Периодичность'
    
    data4 = [
        ['Ежедневный', 'Ежедневно'],
        ['Еженедельный', 'Еженедельно'],
        ['Ежемесячный', 'Ежемесячно'],
        ['Квартальный', 'Ежеквартально'],
        ['Годовой', 'Ежегодно']
    ]
    
    for i, row_data in enumerate(data4, 1):
        row_cells = table4.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Раздел 4
    doc.add_page_break()
    doc.add_heading('4. Технические требования', 1)
    
    doc.add_heading('4.1. Требования к серверному оборудованию', 2)
    
    doc.add_paragraph(
        'Система должна работать на серверах с минимальными характеристиками: '
        'процессор не менее 4 ядер, оперативная память не менее 16 ГБ.'
    )
    
    # Таблица 5 - Требования к оборудованию
    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Light Grid Accent 1'
    header_cells = table5.rows[0].cells
    header_cells[0].text = 'Компонент'
    header_cells[1].text = 'Минимум'
    header_cells[2].text = 'Рекомендуется'
    
    data5 = [
        ['Процессор', '4 ядра', '8 ядер'],
        ['ОЗУ', '16 ГБ', '32 ГБ'],
        ['Диск', '500 ГБ SSD', '1 ТБ SSD'],
        ['Сеть', '1 Гбит/с', '10 Гбит/с'],
        ['ОС', 'Linux/Windows Server', 'Linux/Windows Server']
    ]
    
    for i, row_data in enumerate(data5, 1):
        row_cells = table5.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('4.2. Требования к программному обеспечению', 2)
    
    doc.add_paragraph(
        'Система должна работать на следующих платформах: Windows Server 2016 и выше, '
        'Linux (Ubuntu 18.04 и выше, CentOS 7 и выше).'
    )
    
    # Изображение 4
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 4.1 - Схема развертывания:')
    img4_path = create_simple_image(500, 350, 'lightcoral', 'Развертывание', '../documents/temp_img_ext1_4.png')
    doc.add_picture(img4_path, width=Inches(5))
    
    # Раздел 5
    doc.add_page_break()
    doc.add_heading('5. Требования к интерфейсу', 1)
    
    doc.add_heading('5.1. Общие требования', 2)
    
    doc.add_paragraph(
        'Пользовательский интерфейс должен быть интуитивно понятным и удобным. '
        'Поддержка адаптивного дизайна для различных устройств.'
    )
    
    doc.add_heading('5.2. Требования к дизайну', 2)
    
    doc.add_paragraph(
        'Интерфейс должен соответствовать современным стандартам дизайна. '
        'Поддержка темной и светлой темы оформления.'
    )
    
    # Раздел 6
    doc.add_page_break()
    doc.add_heading('6. Требования к безопасности', 1)
    
    doc.add_heading('6.1. Защита данных', 2)
    
    doc.add_paragraph(
        'Система должна обеспечивать шифрование данных при передаче и хранении. '
        'Использование протоколов TLS 1.2 и выше.'
    )
    
    # Таблица 6 - Уровни безопасности
    table6 = doc.add_table(rows=5, cols=2)
    table6.style = 'Light Grid Accent 1'
    header_cells = table6.rows[0].cells
    header_cells[0].text = 'Уровень'
    header_cells[1].text = 'Описание'
    
    data6 = [
        ['Высокий', 'Шифрование всех данных'],
        ['Средний', 'Шифрование критичных данных'],
        ['Базовый', 'Стандартная защита']
    ]
    
    for i, row_data in enumerate(data6, 1):
        row_cells = table6.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('6.2. Аудит и логирование', 2)
    
    doc.add_paragraph(
        'Система должна вести подробные логи всех операций. '
        'Хранение логов не менее 1 года.'
    )
    
    # Изображение 5
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 6.1 - Схема безопасности:')
    img5_path = create_simple_image(400, 300, 'lightpink', 'Безопасность', '../documents/temp_img_ext1_5.png')
    doc.add_picture(img5_path, width=Inches(4))
    
    # Раздел 7
    doc.add_page_break()
    doc.add_heading('7. Этапы разработки', 1)
    
    doc.add_heading('7.1. Этап 1: Проектирование', 2)
    
    doc.add_paragraph(
        'На этапе проектирования выполняется разработка архитектуры системы, '
        'проектирование базы данных, разработка технического проекта.'
    )
    
    # Таблица 7 - Этапы разработки
    table7 = doc.add_table(rows=6, cols=3)
    table7.style = 'Light Grid Accent 1'
    header_cells = table7.rows[0].cells
    header_cells[0].text = 'Этап'
    header_cells[1].text = 'Срок'
    header_cells[2].text = 'Результат'
    
    data7 = [
        ['Проектирование', '2 месяца', 'Технический проект'],
        ['Разработка', '6 месяцев', 'Рабочая версия'],
        ['Тестирование', '2 месяца', 'Протестированная система'],
        ['Внедрение', '1 месяц', 'Внедренная система'],
        ['Поддержка', '12 месяцев', 'Стабильная работа']
    ]
    
    for i, row_data in enumerate(data7, 1):
        row_cells = table7.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('7.2. Этап 2: Разработка', 2)
    
    doc.add_paragraph(
        'На этапе разработки выполняется программирование модулей системы, '
        'интеграция компонентов, разработка интерфейсов.'
    )
    
    # Изображение 6
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 7.1 - Диаграмма Ганта:')
    img6_path = create_simple_image(600, 200, 'lightcyan', 'Диаграмма Ганта', '../documents/temp_img_ext1_6.png')
    doc.add_picture(img6_path, width=Inches(6))
    
    # Раздел 8
    doc.add_page_break()
    doc.add_heading('8. Заключение', 1)
    
    doc.add_paragraph(
        'Разработка системы должна быть выполнена в соответствии с настоящим '
        'техническим заданием. Срок разработки составляет 12 месяцев с момента '
        'подписания договора.'
    )
    
    doc.add_paragraph(
        'Система должна пройти все этапы тестирования и быть готова к промышленной эксплуатации.'
    )
    
    # Удаление временных файлов
    temp_files = [
        '../documents/temp_img_ext1_1.png', '../documents/temp_img_ext1_2.png',
        '../documents/temp_img_ext1_3.png', '../documents/temp_img_ext1_4.png',
        '../documents/temp_img_ext1_5.png', '../documents/temp_img_ext1_6.png'
    ]
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)
    
    doc.save('../documents/extended_test_document_1.docx')
    print("Создан файл: ../documents/extended_test_document_1.docx")


def create_extended_document_2():
    """Создание второго масштабного тестового документа (с изменениями)."""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Техническое задание на разработку информационной системы управления предприятием', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Содержание (с изменениями)
    doc.add_page_break()
    doc.add_heading('Содержание', 1)
    doc.add_paragraph('1. Общие положения')
    doc.add_paragraph('2. Требования к системе')
    doc.add_paragraph('3. Функциональные требования')
    doc.add_paragraph('4. Технические требования')
    doc.add_paragraph('5. Требования к интерфейсу')
    doc.add_paragraph('6. Требования к безопасности')
    doc.add_paragraph('7. Этапы разработки')
    doc.add_paragraph('8. Интеграция с внешними системами')  # Новый раздел
    doc.add_paragraph('9. Заключение')
    
    # Раздел 1 (с грамматическими изменениями)
    doc.add_page_break()
    doc.add_heading('1. Общие положения', 1)
    
    # Грамматические изменения: добавлена запятая
    doc.add_paragraph(
        'Настоящее техническое задание определяет требования к разработке '
        'информационной системы управления предприятием. Система предназначена '
        'для комплексной автоматизации всех бизнес-процессов организации, и повышения эффективности работы.'
    )
    
    doc.add_heading('1.1. Назначение системы', 2)
    
    doc.add_paragraph(
        'Информационная система управления предприятием предназначена для решения '
        'следующих основных задач:'
    )
    
    doc.add_paragraph('• Автоматизация учета и управления ресурсами предприятия', style='List Bullet')
    doc.add_paragraph('• Управление документооборотом и делопроизводством', style='List Bullet')
    doc.add_paragraph('• Контроль выполнения задач и проектов', style='List Bullet')
    doc.add_paragraph('• Аналитика и формирование отчетности', style='List Bullet')
    doc.add_paragraph('• Интеграция с внешними системами', style='List Bullet')
    doc.add_paragraph('• Управление персоналом и кадровый учет', style='List Bullet')  # Новый пункт
    
    doc.add_heading('1.2. Область применения', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Система предназначена для использования в организациях различных отраслей '
        'экономики. Может быть адаптирована под специфику конкретного предприятия, и интегрирована с существующими системами.'
    )
    
    # Таблица 1 - Области применения (с изменениями)
    doc.add_heading('1.3. Области применения системы', 2)
    
    table1 = doc.add_table(rows=7, cols=2)  # Добавлена строка
    table1.style = 'Light Grid Accent 1'
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Отрасль'
    header_cells[1].text = 'Применение'
    
    data1 = [
        ['Производство', 'Управление производственными процессами и качеством'],
        ['Торговля', 'Управление продажами, складом и логистикой'],  # Изменено
        ['Услуги', 'Управление клиентской базой и сервисами'],
        ['Строительство', 'Управление проектами и ресурсами'],  # Изменено
        ['Образование', 'Управление учебным процессом и студентами'],  # Изменено
        ['Медицина', 'Управление медицинскими записями']  # Новая строка
    ]
    
    for i, row_data in enumerate(data1, 1):
        row_cells = table1.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 1 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 1.1 - Схема архитектуры системы:')
    img1_path = create_simple_image(500, 300, 'lightgreen', 'Архитектура v2.0', '../documents/temp_img_ext2_1.png')  # Изменен цвет и текст
    doc.add_picture(img1_path, width=Inches(5))
    
    # Раздел 2 (с изменениями)
    doc.add_page_break()
    doc.add_heading('2. Требования к системе', 1)
    
    doc.add_heading('2.1. Общие требования', 2)
    
    # Грамматические изменения: изменена запятая
    doc.add_paragraph(
        'Система должна обеспечивать надежную работу в режиме 24/7. '
        'Время отклика системы не должно превышать 1.5 секунд для стандартных операций.'  # Изменено значение
    )
    
    doc.add_heading('2.2. Требования к производительности', 2)
    
    # Изменения: новые значения
    doc.add_paragraph(
        'Система должна обрабатывать не менее 8000 транзакций в минуту. '  # Изменено
        'Поддерживать одновременную работу не менее 1000 пользователей.'  # Изменено
    )
    
    # Таблица 2 - Требования к производительности (с изменениями)
    table2 = doc.add_table(rows=6, cols=3)  # Добавлена строка
    table2.style = 'Light Grid Accent 1'
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Требование'
    header_cells[2].text = 'Единица измерения'
    
    data2 = [
        ['Пропускная способность', '8000', 'транзакций/мин'],  # Изменено
        ['Время отклика', '1.5', 'секунд'],  # Изменено
        ['Количество пользователей', '1000', 'одновременно'],  # Изменено
        ['Доступность', '99.95', '%'],  # Изменено
        ['Версия системы', '2.0.4', '']  # Новая строка
    ]
    
    for i, row_data in enumerate(data2, 1):
        row_cells = table2.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('2.3. Требования к масштабируемости', 2)
    
    # Грамматические изменения: добавлена запятая
    doc.add_paragraph(
        'Система должна поддерживать горизонтальное масштабирование, и вертикальное масштабирование. '
        'Возможность добавления новых серверов без остановки работы системы.'
    )
    
    # Изображение 2 (новое)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 2.1 - Схема масштабирования:')
    img2_path = create_simple_image(400, 250, 'lightblue', 'Масштабирование v2', '../documents/temp_img_ext2_2.png')
    doc.add_picture(img2_path, width=Inches(4))
    
    # Новое изображение
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 2.2 - Схема кластеризации:')
    img2_2_path = create_simple_image(350, 200, 'lightyellow', 'Кластер', '../documents/temp_img_ext2_2_2.png')
    doc.add_picture(img2_2_path, width=Inches(3.5))
    
    # Раздел 3 (с изменениями)
    doc.add_page_break()
    doc.add_heading('3. Функциональные требования', 1)
    
    doc.add_heading('3.1. Модуль управления пользователями', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Модуль должен обеспечивать создание, редактирование, и удаление учетных записей пользователей. '
        'Поддержка ролевой модели доступа, и многофакторной аутентификации.'
    )
    
    doc.add_heading('3.1.1. Управление ролями', 3)
    
    doc.add_paragraph('Система должна поддерживать следующие роли:')
    doc.add_paragraph('• Администратор системы', style='List Bullet')
    doc.add_paragraph('• Менеджер', style='List Bullet')
    doc.add_paragraph('• Пользователь', style='List Bullet')
    doc.add_paragraph('• Гость', style='List Bullet')
    doc.add_paragraph('• Аудитор', style='List Bullet')  # Новый пункт
    
    doc.add_heading('3.1.2. Аутентификация', 3)
    
    # Изменения: добавлен текст
    doc.add_paragraph(
        'Система должна обеспечивать безопасную аутентификацию пользователей. '
        'Поддержка двухфакторной аутентификации, биометрической аутентификации, и интеграции с Active Directory.'
    )
    
    # Таблица 3 - Роли и права доступа (с изменениями)
    table3 = doc.add_table(rows=6, cols=4)  # Добавлена строка
    table3.style = 'Light Grid Accent 1'
    header_cells = table3.rows[0].cells
    header_cells[0].text = 'Роль'
    header_cells[1].text = 'Просмотр'
    header_cells[2].text = 'Редактирование'
    header_cells[3].text = 'Удаление'
    
    data3 = [
        ['Администратор', 'Да', 'Да', 'Да'],
        ['Менеджер', 'Да', 'Да', 'Ограничен'],  # Изменено
        ['Пользователь', 'Да', 'Нет', 'Нет'],
        ['Гость', 'Ограничен', 'Нет', 'Нет'],  # Изменено
        ['Аудитор', 'Да', 'Нет', 'Нет']  # Новая строка
    ]
    
    for i, row_data in enumerate(data3, 1):
        row_cells = table3.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('3.2. Модуль управления документами', 2)
    
    # Грамматические изменения: изменена запятая
    doc.add_paragraph(
        'Модуль обеспечивает полный цикл работы с документами: создание, редактирование, '
        'согласование, утверждение, архивирование, и удаление.'
    )
    
    doc.add_heading('3.2.1. Создание документов', 3)
    
    # Изменения: добавлен текст
    doc.add_paragraph(
        'Система должна поддерживать создание документов различных типов: '
        'приказы, распоряжения, договоры, акты, протоколы, и другие. Поддержка шаблонов документов.'
    )
    
    doc.add_heading('3.2.2. Маршрутизация документов', 3)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Система должна обеспечивать настройку маршрутов движения документов. '
        'Поддержка параллельного, последовательного, и условного согласования.'
    )
    
    # Изображение 3 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 3.1 - Схема маршрутизации документов:')
    img3_path = create_simple_image(450, 300, 'lightcoral', 'Маршрутизация v2', '../documents/temp_img_ext2_3.png')  # Изменен цвет
    doc.add_picture(img3_path, width=Inches(4.5))
    
    doc.add_heading('3.3. Модуль отчетности', 2)
    
    # Грамматические изменения: добавлена запятая
    doc.add_paragraph(
        'Модуль должен обеспечивать формирование различных типов отчетов: '
        'статистические, аналитические, регламентированные, и пользовательские.'
    )
    
    # Таблица 4 - Типы отчетов (с изменениями)
    table4 = doc.add_table(rows=7, cols=2)  # Добавлена строка
    table4.style = 'Light Grid Accent 1'
    header_cells = table4.rows[0].cells
    header_cells[0].text = 'Тип отчета'
    header_cells[1].text = 'Периодичность'
    
    data4 = [
        ['Ежедневный', 'Ежедневно'],
        ['Еженедельный', 'Еженедельно'],
        ['Ежемесячный', 'Ежемесячно'],
        ['Квартальный', 'Ежеквартально'],
        ['Годовой', 'Ежегодно'],
        ['По требованию', 'По запросу']  # Новая строка
    ]
    
    for i, row_data in enumerate(data4, 1):
        row_cells = table4.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Раздел 4 (с изменениями)
    doc.add_page_break()
    doc.add_heading('4. Технические требования', 1)
    
    doc.add_heading('4.1. Требования к серверному оборудованию', 2)
    
    # Изменения: новые значения
    doc.add_paragraph(
        'Система должна работать на серверах с минимальными характеристиками: '
        'процессор не менее 8 ядер, оперативная память не менее 32 ГБ.'  # Изменено
    )
    
    # Таблица 5 - Требования к оборудованию (с изменениями)
    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Light Grid Accent 1'
    header_cells = table5.rows[0].cells
    header_cells[0].text = 'Компонент'
    header_cells[1].text = 'Минимум'
    header_cells[2].text = 'Рекомендуется'
    
    data5 = [
        ['Процессор', '8 ядер', '16 ядер'],  # Изменено
        ['ОЗУ', '32 ГБ', '64 ГБ'],  # Изменено
        ['Диск', '1 ТБ SSD', '2 ТБ SSD'],  # Изменено
        ['Сеть', '10 Гбит/с', '25 Гбит/с'],  # Изменено
        ['ОС', 'Linux/Windows Server', 'Linux/Windows Server']
    ]
    
    for i, row_data in enumerate(data5, 1):
        row_cells = table5.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('4.2. Требования к программному обеспечению', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Система должна работать на следующих платформах: Windows Server 2016 и выше, '
        'Linux (Ubuntu 18.04 и выше, CentOS 7 и выше, Debian 10 и выше).'  # Добавлено
    )
    
    # Изображение 4 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 4.1 - Схема развертывания:')
    img4_path = create_simple_image(500, 350, 'lightpink', 'Развертывание v2', '../documents/temp_img_ext2_4.png')  # Изменен цвет
    doc.add_picture(img4_path, width=Inches(5))
    
    # Раздел 5 (с изменениями)
    doc.add_page_break()
    doc.add_heading('5. Требования к интерфейсу', 1)
    
    doc.add_heading('5.1. Общие требования', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Пользовательский интерфейс должен быть интуитивно понятным, и удобным. '
        'Поддержка адаптивного дизайна для различных устройств, и мобильных приложений.'
    )
    
    doc.add_heading('5.2. Требования к дизайну', 2)
    
    # Изменения: добавлен текст
    doc.add_paragraph(
        'Интерфейс должен соответствовать современным стандартам дизайна. '
        'Поддержка темной и светлой темы оформления. Поддержка кастомизации цветовой схемы.'
    )
    
    # Раздел 6 (с изменениями)
    doc.add_page_break()
    doc.add_heading('6. Требования к безопасности', 1)
    
    doc.add_heading('6.1. Защита данных', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Система должна обеспечивать шифрование данных при передаче, и хранении. '
        'Использование протоколов TLS 1.3 и выше.'  # Изменено
    )
    
    # Таблица 6 - Уровни безопасности (с изменениями)
    table6 = doc.add_table(rows=5, cols=2)
    table6.style = 'Light Grid Accent 1'
    header_cells = table6.rows[0].cells
    header_cells[0].text = 'Уровень'
    header_cells[1].text = 'Описание'
    
    data6 = [
        ['Высокий', 'Шифрование всех данных, двухфакторная аутентификация'],  # Изменено
        ['Средний', 'Шифрование критичных данных, однофакторная аутентификация'],  # Изменено
        ['Базовый', 'Стандартная защита, парольная аутентификация']  # Изменено
    ]
    
    for i, row_data in enumerate(data6, 1):
        row_cells = table6.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('6.2. Аудит и логирование', 2)
    
    # Изменения: изменены значения
    doc.add_paragraph(
        'Система должна вести подробные логи всех операций. '
        'Хранение логов не менее 3 лет.'  # Изменено
    )
    
    # Изображение 5 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 6.1 - Схема безопасности:')
    img5_path = create_simple_image(400, 300, 'lightcyan', 'Безопасность v2', '../documents/temp_img_ext2_5.png')  # Изменен цвет
    doc.add_picture(img5_path, width=Inches(4))
    
    # Раздел 7 (с изменениями)
    doc.add_page_break()
    doc.add_heading('7. Этапы разработки', 1)
    
    doc.add_heading('7.1. Этап 1: Проектирование', 2)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'На этапе проектирования выполняется разработка архитектуры системы, '
        'проектирование базы данных, разработка технического проекта, и создание прототипов интерфейсов.'
    )
    
    # Таблица 7 - Этапы разработки (с изменениями)
    table7 = doc.add_table(rows=6, cols=3)
    table7.style = 'Light Grid Accent 1'
    header_cells = table7.rows[0].cells
    header_cells[0].text = 'Этап'
    header_cells[1].text = 'Срок'
    header_cells[2].text = 'Результат'
    
    data7 = [
        ['Проектирование', '3 месяца', 'Технический проект'],  # Изменено
        ['Разработка', '8 месяцев', 'Рабочая версия'],  # Изменено
        ['Тестирование', '2 месяца', 'Протестированная система'],
        ['Внедрение', '1 месяц', 'Внедренная система'],
        ['Поддержка', '24 месяца', 'Стабильная работа']  # Изменено
    ]
    
    for i, row_data in enumerate(data7, 1):
        row_cells = table7.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('7.2. Этап 2: Разработка', 2)
    
    # Изменения: добавлен текст
    doc.add_paragraph(
        'На этапе разработки выполняется программирование модулей системы, '
        'интеграция компонентов, разработка интерфейсов, и написание документации.'
    )
    
    # Изображение 6 (измененное)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 7.1 - Диаграмма Ганта:')
    img6_path = create_simple_image(600, 200, 'lightsteelblue', 'Диаграмма Ганта v2', '../documents/temp_img_ext2_6.png')  # Изменен цвет
    doc.add_picture(img6_path, width=Inches(6))
    
    # Новый раздел 8
    doc.add_page_break()
    doc.add_heading('8. Интеграция с внешними системами', 1)
    
    doc.add_heading('8.1. Типы интеграций', 2)
    
    doc.add_paragraph(
        'Система должна поддерживать интеграцию с различными внешними системами: '
        '1С, SAP, Oracle, Microsoft Dynamics, и другими.'
    )
    
    # Таблица 8 - Интеграции (новая таблица)
    table8 = doc.add_table(rows=5, cols=2)
    table8.style = 'Light Grid Accent 1'
    header_cells = table8.rows[0].cells
    header_cells[0].text = 'Система'
    header_cells[1].text = 'Тип интеграции'
    
    data8 = [
        ['1С', 'API, файловый обмен'],
        ['SAP', 'RFC, IDoc'],
        ['Oracle', 'API, база данных'],
        ['Microsoft Dynamics', 'API, веб-сервисы']
    ]
    
    for i, row_data in enumerate(data8, 1):
        row_cells = table8.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение 7 (новое)
    doc.add_paragraph()
    doc.add_paragraph('Рисунок 8.1 - Схема интеграций:')
    img7_path = create_simple_image(500, 300, 'lightgoldenrodyellow', 'Интеграции', '../documents/temp_img_ext2_7.png')
    doc.add_picture(img7_path, width=Inches(5))
    
    # Раздел 9 (было 8)
    doc.add_page_break()
    doc.add_heading('9. Заключение', 1)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Разработка системы должна быть выполнена в соответствии с настоящим '
        'техническим заданием. Срок разработки составляет 14 месяцев с момента '
        'подписания договора.'  # Изменено
    )
    
    doc.add_paragraph(
        'Система должна пройти все этапы тестирования, и быть готова к промышленной эксплуатации.'
    )
    
    # Удаление временных файлов
    temp_files = [
        '../documents/temp_img_ext2_1.png', '../documents/temp_img_ext2_2.png',
        '../documents/temp_img_ext2_2_2.png', '../documents/temp_img_ext2_3.png',
        '../documents/temp_img_ext2_4.png', '../documents/temp_img_ext2_5.png',
        '../documents/temp_img_ext2_6.png', '../documents/temp_img_ext2_7.png'
    ]
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)
    
    doc.save('../documents/extended_test_document_2.docx')
    print("Создан файл: ../documents/extended_test_document_2.docx")


if __name__ == "__main__":
    print("Создание масштабных тестовых документов...")
    print("=" * 60)
    
    create_extended_document_1()
    create_extended_document_2()
    
    print("=" * 60)
    print("Все масштабные тестовые документы созданы успешно!")
    print("\nСозданные файлы:")
    print("  - ../documents/extended_test_document_1.docx (базовый документ)")
    print("  - ../documents/extended_test_document_2.docx (версия с изменениями)")
    print("\nХарактеристики документов:")
    print("  - 8-9 разделов")
    print("  - Множество подразделов и пунктов")
    print("  - 7-8 таблиц в каждом документе")
    print("  - 6-8 изображений в каждом документе")
    print("  - Грамматические изменения (запятые, пунктуация)")
    print("  - Изменения в таблицах и изображениях")

