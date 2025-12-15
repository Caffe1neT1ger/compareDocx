"""
Скрипт для создания тестовых DOCX документов.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import os
import io


def create_test_document_1():
    """Создание первого тестового документа."""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Техническое задание на разработку системы', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1
    doc.add_heading('1. Общие положения', 1)
    
    doc.add_paragraph(
        'Настоящее техническое задание определяет требования к разработке '
        'информационной системы управления документооборотом.'
    )
    
    doc.add_paragraph(
        'Система предназначена для автоматизации процессов работы с документами '
        'в организации и повышения эффективности управления.'
    )
    
    # Подраздел
    doc.add_heading('1.1. Назначение системы', 2)
    
    doc.add_paragraph(
        'Информационная система предназначена для решения следующих задач:'
    )
    
    doc.add_paragraph('• Учет и регистрация входящих и исходящих документов', style='List Bullet')
    doc.add_paragraph('• Контроль исполнения документов', style='List Bullet')
    doc.add_paragraph('• Поиск документов по различным критериям', style='List Bullet')
    doc.add_paragraph('• Формирование отчетов и статистики', style='List Bullet')
    
    # Раздел 2
    doc.add_heading('2. Требования к функциональным характеристикам', 1)
    
    doc.add_paragraph(
        'Система должна обеспечивать выполнение следующих функций:'
    )
    
    # Глава 2.1
    doc.add_heading('2.1. Управление документами', 2)
    
    doc.add_paragraph(
        'Модуль управления документами должен предоставлять возможность:'
    )
    
    doc.add_paragraph(
        'Создания новых документов с указанием всех необходимых реквизитов, '
        'включая дату создания, автора, тип документа и статус.'
    )
    
    doc.add_paragraph(
        'Редактирования существующих документов с сохранением истории изменений.'
    )
    
    # Глава 2.2
    doc.add_heading('2.2. Поиск и фильтрация', 2)
    
    doc.add_paragraph(
        'Система должна обеспечивать поиск документов по следующим параметрам:'
    )
    
    doc.add_paragraph('• По названию или содержимому', style='List Bullet')
    doc.add_paragraph('• По дате создания или изменения', style='List Bullet')
    doc.add_paragraph('• По автору документа', style='List Bullet')
    doc.add_paragraph('• По типу документа', style='List Bullet')
    
    # Раздел 3
    doc.add_heading('3. Требования к интерфейсу', 1)
    
    doc.add_paragraph(
        'Пользовательский интерфейс должен быть интуитивно понятным и удобным '
        'для работы. Все основные функции должны быть доступны не более чем '
        'за два клика мыши.'
    )
    
    doc.add_paragraph(
        'Интерфейс должен поддерживать работу с различными разрешениями экрана '
        'и адаптироваться под размер окна браузера.'
    )
    
    # Таблица с характеристиками
    doc.add_heading('3.1. Технические характеристики', 2)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Требование'
    header_cells[2].text = 'Примечание'
    
    # Данные таблицы
    data_rows = [
        ['Производительность', 'Обработка не менее 1000 документов/час', 'Минимальные требования'],
        ['Хранилище', 'Не менее 1 ТБ', 'С возможностью расширения'],
        ['Доступность', '99.9%', 'В рабочее время']
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Добавление изображения (создаем простое тестовое изображение)
    doc.add_paragraph()  # Пустая строка
    doc.add_paragraph('Схема архитектуры системы:')
    
    # Создаем простое изображение
    img = Image.new('RGB', (400, 200), color='lightblue')
    draw = ImageDraw.Draw(img)
    draw.rectangle([50, 50, 350, 150], outline='black', width=2)
    draw.text((150, 90), 'Система управления', fill='black')
    draw.text((150, 110), 'документооборотом', fill='black')
    
    # Сохраняем изображение во временный файл
    img_path = '../documents/temp_img_1.png'
    img.save(img_path)
    
    # Добавляем изображение в документ
    doc.add_picture(img_path, width=Inches(4))
    
    # Заключение
    doc.add_heading('Заключение', 1)
    
    doc.add_paragraph(
        'Разработка системы должна быть выполнена в соответствии с настоящим '
        'техническим заданием и действующими стандартами в области разработки '
        'программного обеспечения.'
    )
    
    # Удаляем временный файл изображения
    if os.path.exists(img_path):
        os.remove(img_path)
    
    doc.save('../documents/test_document_1.docx')
    print("Создан файл: ../documents/test_document_1.docx")


def create_test_document_2():
    """Создание второго тестового документа (с изменениями)."""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Техническое задание на разработку системы', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1
    doc.add_heading('1. Общие положения', 1)
    
    # Грамматические изменения: добавлена запятая, изменено слово
    doc.add_paragraph(
        'Настоящее техническое задание определяет требования к разработке '
        'информационной системы управления документооборотом, и архивацией документов.'
    )
    
    # Грамматические изменения: изменен порядок слов, добавлена запятая
    doc.add_paragraph(
        'Система предназначена для автоматизации процессов работы с документами '
        'в организации, для повышения эффективности управления, и обеспечения безопасности данных.'
    )
    
    # Подраздел
    doc.add_heading('1.1. Назначение системы', 2)
    
    doc.add_paragraph(
        'Информационная система предназначена для решения следующих задач:'
    )
    
    doc.add_paragraph('• Учет и регистрация входящих и исходящих документов', style='List Bullet')
    doc.add_paragraph('• Контроль исполнения документов', style='List Bullet')
    doc.add_paragraph('• Поиск документов по различным критериям', style='List Bullet')
    doc.add_paragraph('• Формирование отчетов и статистики', style='List Bullet')
    doc.add_paragraph('• Архивирование документов', style='List Bullet')
    
    # Раздел 2
    doc.add_heading('2. Требования к функциональным характеристикам', 1)
    
    doc.add_paragraph(
        'Система должна обеспечивать выполнение следующих функций:'
    )
    
    # Глава 2.1
    doc.add_heading('2.1. Управление документами', 2)
    
    doc.add_paragraph(
        'Модуль управления документами должен предоставлять возможность:'
    )
    
    # Грамматические изменения: изменена запятая, добавлено слово
    doc.add_paragraph(
        'Создания новых документов с указанием всех необходимых реквизитов '
        'включая дату создания, автора, тип документа, статус, и категорию.'
    )
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Редактирования существующих документов, с сохранением полной истории изменений, '
        'и возможностью отката к предыдущим версиям.'
    )
    
    doc.add_paragraph(
        'Удаления документов с возможностью восстановления в течение определенного периода.'
    )
    
    # Глава 2.2
    doc.add_heading('2.2. Поиск и фильтрация', 2)
    
    # Грамматические изменения: изменено слово
    doc.add_paragraph(
        'Система должна обеспечивать расширенный поиск документов по следующим параметрам:'
    )
    
    doc.add_paragraph('• По названию или содержимому (полнотекстовый поиск)', style='List Bullet')
    doc.add_paragraph('• По дате создания или изменения', style='List Bullet')
    doc.add_paragraph('• По автору документа', style='List Bullet')
    doc.add_paragraph('• По типу документа', style='List Bullet')
    doc.add_paragraph('• По тегам и метаданным', style='List Bullet')
    
    # Новый раздел
    doc.add_heading('2.3. Безопасность данных', 2)
    
    doc.add_paragraph(
        'Система должна обеспечивать защиту данных на всех уровнях:'
    )
    
    doc.add_paragraph('• Шифрование данных при хранении', style='List Bullet')
    doc.add_paragraph('• Контроль доступа на основе ролей', style='List Bullet')
    doc.add_paragraph('• Аудит всех операций с документами', style='List Bullet')
    
    # Раздел 3
    doc.add_heading('3. Требования к интерфейсу', 1)
    
    # Грамматические изменения: изменена пунктуация
    doc.add_paragraph(
        'Пользовательский интерфейс должен быть интуитивно понятным современным '
        'и удобным для работы. Все основные функции должны быть доступны, не более чем '
        'за два клика мыши.'
    )
    
    doc.add_paragraph(
        'Интерфейс должен поддерживать работу с различными разрешениями экрана '
        'адаптироваться под размер окна браузера, и поддерживать мобильные устройства.'
    )
    
    doc.add_paragraph(
        'Необходимо обеспечить поддержку темной, и светлой темы оформления.'
    )
    
    # Таблица с изменениями (добавлена строка, изменены данные)
    doc.add_heading('3.1. Технические характеристики', 2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Требование'
    header_cells[2].text = 'Примечание'
    
    # Данные таблицы (с изменениями)
    data_rows = [
        ['Производительность', 'Обработка не менее 1500 документов/час', 'Обновленные требования'],
        ['Хранилище', 'Не менее 2 ТБ', 'С возможностью расширения'],
        ['Доступность', '99.95%', 'В рабочее время'],
        ['Версия системы', '2.0.4', 'Текущая версия']  # Новая строка
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Измененное изображение
    doc.add_paragraph()  # Пустая строка
    doc.add_paragraph('Схема архитектуры системы:')
    
    # Создаем измененное изображение
    img = Image.new('RGB', (400, 200), color='lightgreen')  # Изменен цвет
    draw = ImageDraw.Draw(img)
    draw.rectangle([50, 50, 350, 150], outline='black', width=2)
    draw.text((150, 90), 'Система управления', fill='black')
    draw.text((150, 110), 'документооборотом v2.0', fill='black')  # Добавлен текст версии
    
    # Сохраняем изображение во временный файл
    img_path = '../documents/temp_img_2.png'
    img.save(img_path)
    
    # Добавляем изображение в документ
    doc.add_picture(img_path, width=Inches(4))
    
    # Добавляем еще одно изображение (новое)
    doc.add_paragraph()  # Пустая строка
    doc.add_paragraph('Диаграмма процессов:')
    
    img2 = Image.new('RGB', (300, 150), color='lightyellow')
    draw2 = ImageDraw.Draw(img2)
    draw2.ellipse([50, 30, 250, 120], outline='black', width=2)
    draw2.text((100, 65), 'Процесс', fill='black')
    
    img2_path = '../documents/temp_img_2_2.png'
    img2.save(img2_path)
    doc.add_picture(img2_path, width=Inches(3))
    
    # Заключение
    doc.add_heading('Заключение', 1)
    
    doc.add_paragraph(
        'Разработка системы должна быть выполнена в соответствии с настоящим '
        'техническим заданием, действующими стандартами в области разработки '
        'программного обеспечения и требованиями информационной безопасности.'
    )
    
    doc.add_paragraph(
        'Срок разработки системы составляет 6 месяцев с момента подписания договора.'
    )
    
    # Удаляем временные файлы изображений
    for img_file in [img_path, img2_path]:
        if os.path.exists(img_file):
            os.remove(img_file)
    
    doc.save('../documents/test_document_2.docx')
    print("Создан файл: ../documents/test_document_2.docx")


def create_test_document_3():
    """Создание третьего тестового документа (совсем другой)."""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчет о проделанной работе', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Введение
    doc.add_heading('Введение', 1)
    
    doc.add_paragraph(
        'Настоящий отчет содержит информацию о проделанной работе за отчетный период.'
    )
    
    # Раздел 1
    doc.add_heading('1. Выполненные задачи', 1)
    
    doc.add_paragraph('В течение отчетного периода были выполнены следующие задачи:')
    
    doc.add_paragraph('• Разработка архитектуры системы', style='List Bullet')
    doc.add_paragraph('• Создание базы данных', style='List Bullet')
    doc.add_paragraph('• Реализация основного функционала', style='List Bullet')
    
    # Раздел 2
    doc.add_heading('2. Результаты', 1)
    
    doc.add_paragraph(
        'В результате проделанной работы была создана рабочая версия системы, '
        'прошедшая первичное тестирование.'
    )
    
    # Таблица с результатами
    doc.add_heading('2.1. Статистика разработки', 2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Показатель'
    header_cells[1].text = 'Значение'
    
    data_rows = [
        ['Модулей разработано', '15'],
        ['Строк кода', '45000'],
        ['Тестов написано', '320']
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Изображение
    doc.add_paragraph()  # Пустая строка
    doc.add_paragraph('График выполнения работ:')
    
    img = Image.new('RGB', (350, 200), color='lightcoral')
    draw = ImageDraw.Draw(img)
    # Простой график
    points = [(50, 150), (100, 120), (150, 100), (200, 80), (250, 60), (300, 50)]
    for i in range(len(points) - 1):
        draw.line([points[i], points[i+1]], fill='black', width=2)
    draw.text((100, 160), 'Прогресс разработки', fill='black')
    
    img_path = '../documents/temp_img_3.png'
    img.save(img_path)
    doc.add_picture(img_path, width=Inches(3.5))
    
    # Заключение
    doc.add_heading('Заключение', 1)
    
    doc.add_paragraph('Работа выполнена в полном объеме согласно плану.')
    
    # Удаляем временный файл изображения
    if os.path.exists(img_path):
        os.remove(img_path)
    
    doc.save('../documents/test_document_3.docx')
    print("Создан файл: ../documents/test_document_3.docx")


if __name__ == "__main__":
    print("Создание тестовых документов...")
    print("-" * 50)
    
    create_test_document_1()
    create_test_document_2()
    create_test_document_3()
    
    print("-" * 50)
    print("Все тестовые документы созданы успешно!")
    print("\nСозданные файлы:")
    print("  - ../documents/test_document_1.docx (базовый документ)")
    print("  - ../documents/test_document_2.docx (версия с изменениями)")
    print("  - ../documents/test_document_3.docx (совсем другой документ)")

